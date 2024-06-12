from docx import Document


def read_docx(path: str) -> str:
    doc = Document(path)
    full_text = []

    # Read all paragraphs
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)

    # Read all tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)

    return '\n'.join(full_text)
